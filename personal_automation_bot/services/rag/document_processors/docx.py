"""
Processor for Microsoft Word (DOCX) files.
"""
import os
import logging
from typing import Dict, Any, List, Optional
import re

from personal_automation_bot.services.rag.document_processors.base import DocumentProcessor

logger = logging.getLogger(__name__)

class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word (DOCX) files."""

    def __init__(self):
        """Initialize the DOCX processor."""
        super().__init__()
        self.extensions = ['.docx']

        # Check if required libraries are available
        self.has_python_docx = self._check_python_docx()

        if not self.has_python_docx:
            logger.warning("python-docx is not available. DOCX processing will be limited.")

    def _check_python_docx(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            return False

    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the given file."""
        _, ext = os.path.splitext(file_path.lower())
        return ext in self.extensions and self.has_python_docx

    def extract_text(self, file_path: str, **kwargs) -> str:
        """Extract text from a DOCX file."""
        if not self.has_python_docx:
            raise ImportError("python-docx is not installed. Install it with 'pip install python-docx'")

        import docx

        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))

        # Join all text with newlines
        text = "\\n".join(full_text)

        # Clean up text
        text = self._clean_text(text)

        return text

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Replace multiple newlines with a single one
        text = re.sub(r'\\n{3,}', '\\n\\n', text)

        # Remove excessive whitespace
        text = re.sub(r' {2,}', ' ', text)

        return text.strip()

    def extract_metadata(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Extract metadata from a DOCX file."""
        metadata = {
            "filename": os.path.basename(file_path),
            "extension": ".docx",
            "file_type": "docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }

        # Add file stats
        stat = os.stat(file_path)
        metadata.update({
            "size_bytes": stat.st_size,
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime
        })

        # Try to extract DOCX-specific metadata
        if self.has_python_docx:
            try:
                import docx
                doc = docx.Document(file_path)

                # Extract core properties if available
                if hasattr(doc, 'core_properties'):
                    props = doc.core_properties
                    if props.author:
                        metadata["author"] = props.author
                    if props.title:
                        metadata["title"] = props.title
                    if props.subject:
                        metadata["subject"] = props.subject
                    if props.keywords:
                        metadata["keywords"] = props.keywords
                    if props.created:
                        metadata["doc_created"] = props.created.isoformat()
                    if props.modified:
                        metadata["doc_modified"] = props.modified.isoformat()

                # Count paragraphs, tables, etc.
                metadata["paragraph_count"] = len(doc.paragraphs)
                metadata["table_count"] = len(doc.tables)
            except Exception as e:
                logger.warning(f"Error extracting DOCX metadata: {e}")

        return metadata
